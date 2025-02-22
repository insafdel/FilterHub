from docx import Document
import csv

# Load the Word document
doc = Document("liste_additive_FR[1].docx")  # Replace with actual file name

# Initialize list to store extracted data
data = []

# Loop through all tables in the document
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]  # Get text from each cell

        if len(cells) >= 2:  # Ensure at least Job Code and Job Title are present
            job_code = cells[0]
            job_title = cells[1]

            # If there's a third column, it might contain additional details
            if len(cells) > 2:
                job_title += " - " + cells[2]  # Append extra info if needed

            # Append extracted data
            data.append([job_code, job_title])

# Save to CSV
csv_filename = "jobs_outputfr.csv"
with open(csv_filename, mode="w", newline="", encoding="utf-8") as file:
    writer = csv.writer(file)
    writer.writerow(["Job Code", "Job Title"])  # Write header
    writer.writerows(data)  # Write extracted data

print(f"✅ Word file successfully converted to {csv_filename}")

import csv

# Load the Word document
doc = Document("liste_additive_AR[1].docx")  # Replace with actual file name

# Initialize list to store extracted data
data = []

# Loop through all tables in the document
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]  # Get text from each cell

        if len(cells) >= 2:  # Ensure at least Job Code and Job Title are present
            job_code = cells[1]
            job_title = cells[0]

            # If there's a third column, it might contain additional details
            if len(cells) > 2:
                job_title += "" + cells[2]  # Append extra info if needed

            # Append extracted data
            data.append([job_code, job_title])

# Save to CSV
csv_filename = "jobs_outputar.csv"
with open(csv_filename, mode="w", newline="", encoding="utf-8") as file:
    writer = csv.writer(file)
    writer.writerow(["Job Code", "Job Title"])  # Write header
    writer.writerows(data)  # Write extracted data

print(f"✅ Word file successfully converted to {csv_filename}")

import pandas as pd

# Load French and Arabic CSV files
df_fr = pd.read_csv("jobs_outputar.csv", encoding="utf-8-sig")
df_ar = pd.read_csv("jobs_outputfr.csv", encoding="utf-8-sig")

# Merge both dataframes on 'Job Code'
df_merged = pd.merge(df_fr, df_ar, on="Job Code", suffixes=("_FR", "_AR"))

# Save to new CSV file
df_merged.to_csv("merged_jobs.csv", index=False, encoding="utf-8-sig")

print("✅ Merging complete! Saved as 'merged_jobs.csv'")
df = pd.read_csv("merged_jobs.csv")

df.rename(columns={"Job Title_FR": "ar_name_activity"}, inplace=True)
df.rename(columns={"Job Title_AR": "name_activity"}, inplace=True)

df.rename(columns={"Job Code": "code_activity"}, inplace=True)
df = df[["code_activity","name_activity", "ar_name_activity"]] 
# Load both CSV files
df1 = pd.read_excel("Activity-2025-02-18(1).xlsx")
df2 = df

# Concatenate them (stack rows)
df_merged = pd.concat([df1, df2], ignore_index=True)

# Save the merged file
df_merged.to_csv("merged_data.csv", index=False, encoding="utf-8-sig")

print("✅ CSV files merged successfully! Saved as 'merged_data.csv'")




# __________________________________________________________________the model___________________________________________


import pandas as pd
from sentence_transformers import SentenceTransformer, util

# Load pre-trained multilingual SBERT model
model = SentenceTransformer('paraphrase-multilingual-mpnet-base-v2')

def preprocess_text(text):
    """Basic text preprocessing"""
    if pd.isna(text):
        return ""
    return str(text).lower().strip()

def get_best_match(input_embedding, reference_embeddings, reference_titles, threshold):
    """Get best matching title from embeddings"""
    similarities = util.pytorch_cos_sim(input_embedding, reference_embeddings)[0]
    best_match_idx = similarities.argmax().item()
    best_match_score = similarities[best_match_idx].item()

    return {
        'best_match': reference_titles[best_match_idx] if best_match_score >= threshold else "Other",
        'score': best_match_score,
        'exists': best_match_score >= threshold
    }

def create_job_matcher(excel_file, excel_file02, csv_file, csv_file02, threshold):
    """Match job titles with different thresholds"""

    # Load data
    df_activities = pd.read_excel(excel_file)
    df_rejected_activities_1 = pd.read_excel(excel_file02, header=None)  # No header for commerce file
    df_rejected_activities_2 = pd.read_csv(csv_file02, header=None)  # No header for bita9a file
    df_user_inputs = pd.read_excel(csv_file)

    # Prepare reference data (existing jobs)
    french_titles = df_user_inputs["name_activity"].dropna().tolist()
    arabic_titles = df_user_inputs["ar_name_activity"].dropna().tolist()

    # Prepare rejected activities data
    arabic_rejected_titles_1 = df_rejected_activities_1[0].dropna().tolist()
    arabic_rejected_titles_2 = df_rejected_activities_2[0].dropna().tolist()
    arabic_rejected_titles = arabic_rejected_titles_1 + arabic_rejected_titles_2

    # Combine all reference titles (existing + rejected)
    all_existing_titles = french_titles + arabic_titles
    all_existing_titles = [preprocess_text(title) for title in all_existing_titles]
    all_existing_titles = list(filter(None, all_existing_titles))  # Remove empty strings

    all_rejected_titles = [preprocess_text(title) for title in arabic_rejected_titles]
    all_rejected_titles = list(filter(None, all_rejected_titles))  # Remove empty strings

    # Combine existing and rejected titles into one list
    all_reference_titles = all_existing_titles + all_rejected_titles

    # Create embeddings for all reference titles
    reference_embeddings = model.encode(all_reference_titles)

    results = []

    # Process each new activity
    for _, row in df_activities.iterrows():
        activity = preprocess_text(row['activity'])
        description = preprocess_text(row.get('description', ''))

        if not activity:
            continue

        # Create embedding for the activity
        activity_embedding = model.encode(activity)

        # Get best match based on activity title
        title_match = get_best_match(
            activity_embedding,
            reference_embeddings,
            all_reference_titles,
            threshold
        )

        # Determine category based on the best match
        if title_match['exists']:
            # Check if the best match is in the rejected list
            if title_match['best_match'] in all_rejected_titles:
                category = "rejected"
            else:
                category = "accepted"  # Best match is in existing activities
        else:
            category = "other"  # No match found or below threshold

        results.append({
            'activity_name': activity,
            'description': description,
            'best_match': title_match['best_match'],
            'confidence_score': round(title_match['score'], 3),
            'category': category
        })

    # Convert results to DataFrame
    df_results = pd.DataFrame(results)

    return df_results

# 🔥 **Find the Optimal Threshold**
best_threshold = 0
best_accuracy = 0

# Load the ground truth file
df_ground_truth = pd.read_csv("transformed2_data.csv")
df_ground_truth.columns = df_ground_truth.columns.str.strip()  # Remove extra spaces

# Test different thresholds
for threshold in [0.4, 0.45, 0.5, 0.55, 0.6, 0.65, 0.7, 0.75, 0.8]:
    print(f"🔎 Testing threshold: {threshold}")

    # Run the matcher with the current threshold
    df_results = create_job_matcher(
        excel_file="3000-proposition mars 2024.xlsx",
        excel_file02="Commerce_no_numbers.xlsx",
        csv_file="Activity-2025-02-18(1).xlsx",
        csv_file02="Bita9a_Mihanya (2).csv",
        threshold=threshold
    )

    # Ensure both DataFrames have the same number of rows
    if len(df_results) != len(df_ground_truth):
        print(f"⚠️ Mismatch in row count! Results: {len(df_results)}, Ground Truth: {len(df_ground_truth)}")
        df_ground_truth = df_ground_truth.iloc[:len(df_results)]  # Trim if necessary

    # Reset index to align rows
    df_results = df_results.reset_index(drop=True)
    df_ground_truth = df_ground_truth.reset_index(drop=True)

    # Compare predicted categories with ground truth

    df_results = df_
    results.iloc[:-1]
    df_ground_truth['match'] = df_results['category'] == df_ground_truth['proposition']

    # Calculate accuracy
    accuracy = df_ground_truth['match'].mean() * 100  # Convert to percentage
    print(f"✅ Accuracy for threshold {threshold}: {accuracy:.2f}%")

    # Store the best threshold
    if accuracy > best_accuracy:
        best_accuracy = accuracy
        best_threshold = threshold

# 🎯 **Print the Best Threshold**
print(f"\n🏆 Best Threshold: {best_threshold} with Accuracy: {best_accuracy:.2f}%")



